import os
import docx
# 合并后的文档路径
foutname = 'merge.docx'

# 获取分离的文档
index = 0x00
fname = 'split_' + '{0:02x}'.format(0) + '.docx'

new_word = docx.Document()
word_context = []
while os.path.exists(fname):
    print(fname)
    file = docx.Document(fname)
    # 可惜这种方式，仅仅可以复制过来text内容，
    # 表格、字体格式等均未能保留
    for page in file.paragraphs:
        word_context.append(page.text)
        new_word.add_paragraph(page.text)
    index += 1
    fname = 'split_' + '{0:02x}'.format(index) + '.docx'
print('word merge finished!')
print(word_context)
new_word.save(foutname)
